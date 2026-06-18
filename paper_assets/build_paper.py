# -*- coding: utf-8 -*-
"""Build the CR-HKGE IEEE two-column conference paper (.docx)."""
import os
from docx import Document
from docx.shared import Pt, Mm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r"D:\Azam\MTI\TESIS\knowledge_graph_attention_network"
ASSETS = os.path.join(ROOT, "paper_assets")
OUT = os.path.join(ROOT, "CR-HKGE_IEEE_Conference_Paper.docx")

FONT = "Times New Roman"


# ---------------- low-level helpers ----------------
def set_cols(section, num, space_mm=4.22):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(int(space_mm * 56.7)))  # mm -> twips


def set_margins(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(19)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(14.32)
    section.right_margin = Mm(14.32)


def _runs_font(p, size, bold=False, italic=False):
    for r in p.runs:
        r.font.name = FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        rpr = r._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), FONT)
        rfonts.set(qn("w:hAnsi"), FONT)


def para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=3, line=1.0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    return p


def add_run(p, text, size=10, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)
    return r


def body(doc, text, size=10):
    p = para(doc, after=3)
    add_run(p, text, size=size)
    return p


def h1(doc, text):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=3)
    add_run(p, text, size=10, bold=True)
    return p


def h2(doc, text):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=2)
    add_run(p, text, size=10, bold=True, italic=True)
    return p


def equation(doc, text, num):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=4)
    add_run(p, text, size=10, italic=True)
    add_run(p, "      (" + str(num) + ")", size=10)
    return p


def figure(doc, path, width_in, caption_bold, caption_rest):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=2)
    p.add_run().add_picture(path, width=Inches(width_in))
    c = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=8)
    add_run(c, caption_bold, size=8, bold=False)
    add_run(c, caption_rest, size=8)


def _set_cell(cell, text, size=8, bold=False, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if align == "right"
                   else WD_ALIGN_PARAGRAPH.CENTER if align == "center"
                   else WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    add_run(p, text, size=size, bold=bold)


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def table_caption(doc, num_label, title):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=1)
    add_run(p, num_label, size=8)
    p2 = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2)
    add_run(p2, title, size=8)


def make_table(doc, header, rows, aligns=None, size=8, highlight_last=False, col_w=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.autofit = True
    hdr = t.rows[0].cells
    for i, htext in enumerate(header):
        _set_cell(hdr[i], htext, size=size, bold=True,
                  align="left" if i == 0 else "center")
        _shade(hdr[i], "E8E8E8")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            a = "left" if ci == 0 else (aligns[ci] if aligns else "right")
            _set_cell(cells[ci], val, size=size,
                      bold=(highlight_last and ri == len(rows) - 1),
                      align=a)
        if highlight_last and ri == len(rows) - 1:
            for c in cells:
                _shade(c, "F0F0F0")
    # borders: top/bottom + under header
    _table_borders(t)
    if col_w:
        for row in t.rows:
            for ci, w in enumerate(col_w):
                row.cells[ci].width = Inches(w)
    para(doc, before=2, after=4)
    return t


def _table_borders(t):
    tbl = t._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    for edge in ("left", "right", "insideV"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


# ================= BUILD =================
doc = Document()
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10)

sec0 = doc.sections[0]
set_margins(sec0)
set_cols(sec0, 1)

# ---- Title block (single column) ----
p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=6)
add_run(p, "CR-HKGE: Cross-Reference Semantic Enrichment on "
           "Heterogeneous Knowledge Graph Embedding for "
           "Fragrance Recommendation", size=20, bold=True)

p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=2)
add_run(p, "Azam Azri Ahmad", size=11)
add_run(p, "¹", size=8)
add_run(p, ", Raissa Desyandita", size=11)
add_run(p, "²", size=8)

p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
add_run(p, "Department / Faculty [VERIFY], Institution [VERIFY]", size=10)
p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=10)
add_run(p, "{author1, author2}@institution [VERIFY]", size=10, italic=True)

# ---- switch to two columns ----
doc.add_section(WD_SECTION.CONTINUOUS)
sec1 = doc.sections[1]
set_margins(sec1)
set_cols(sec1, 2)

# ---- Abstract ----
p = para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=4)
add_run(p, "Abstract—", size=9, bold=True, italic=True)
add_run(p,
    "Fragrance recommendation is a challenging cold-start problem because "
    "small-scale perfume catalogs often lack historical user-item interactions, "
    "while product relevance is expressed through subjective olfactory concepts "
    "such as accords, fragrance families, visual notes, and global fragrance "
    "references. Existing knowledge graph (KG) recommendation models, including "
    "KGAT, commonly rely on user-item interaction matrices and do not explicitly "
    "model cross-reference semantics between local perfume products and global "
    "fragrance references. This paper proposes CR-HKGE, a Cross-Reference "
    "Semantic Enrichment model built on heterogeneous knowledge graph embedding "
    "for fragrance recommendation without historical user behavior. CR-HKGE "
    "constructs a fragrance-specific heterogeneous KG, propagates global "
    "reference semantics through the inspired_by relation, and applies "
    "relation-type attention to prioritize informative fragrance relations. Since "
    "the Aromatique dataset contains no purchase, rating, or click logs, "
    "content-based positive pairs are constructed as surrogate supervision for "
    "KGAT-compatible offline training and evaluation in a zero historical "
    "interaction setting. Experiments on 340 perfume products, 998 entities, "
    "9,250 triples, and seven relation types show that CR-HKGE outperforms KGAT "
    "and other KG/collaborative-filtering baselines across overall Top-K metrics, "
    "with the largest gains on enriched products linked to global references. "
    "These results indicate that cross-reference semantic enrichment can improve "
    "KG-based fragrance recommendation when no interaction history is available.",
    size=9)

p = para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=2, after=8)
add_run(p, "Keywords—", size=9, bold=True, italic=True)
add_run(p, "knowledge graph recommendation; heterogeneous knowledge graph; "
           "fragrance recommendation; cold-start recommendation; cross-reference "
           "enrichment; relation-type attention; KGAT", size=9)

# ================= I. INTRODUCTION =================
h1(doc, "I.\tINTRODUCTION")
body(doc,
    "Unlike movie or book recommendation, fragrance recommendation often relies "
    "on subjective olfactory descriptors and reference scents rather than "
    "explicit ratings or purchase histories. Users typically describe what they "
    "want through perceptual concepts such as sweet, woody, amber, or vanilla, "
    "and frequently mention a reference perfume whose scent they wish to "
    "approximate. For a small local catalog, this perceptual and reference-driven "
    "nature makes conventional interaction-based recommendation difficult to "
    "apply directly.")
body(doc,
    "Collaborative filtering and many KG-based recommenders require a user-item "
    "interaction matrix [1], [3], [9]. In the Aromatique catalog studied here, no "
    "historical purchase, rating, or click logs are available. This makes "
    "standard collaborative filtering and KGAT-style training [1] infeasible "
    "without constructing a surrogate supervision signal, since these models "
    "cannot be executed without an interaction source.")
body(doc,
    "Fragrance data are, however, naturally relational. A product is connected to "
    "its accords, visual notes, and olfactory family, and an inspired product is "
    "additionally connected to a global fragrance reference that carries its own "
    "accords and family. The inspired_by relation therefore acts as a semantic "
    "bridge from the local product space to a richer global fragrance knowledge "
    "space, which is well suited to a knowledge graph formulation [7], [10].")
body(doc,
    "Nevertheless, two gaps remain. First, KGAT applies attentive propagation but "
    "does not assign explicit, type-specific priors to domain relations, so a "
    "highly informative relation such as inspired_by is treated similarly to a "
    "weak relation such as has_visual_note [1], [16]. Second, existing KG "
    "recommenders are typically validated on large datasets with abundant user "
    "interactions [6], [14], and prior work reports instability on extremely "
    "sparse data [14], [15]. A small interaction-free fragrance catalog is an "
    "extreme case that these models do not directly address.")
body(doc,
    "To address these gaps, we propose CR-HKGE, a Cross-Reference Heterogeneous "
    "Knowledge Graph Embedding model that modifies KGAT for fragrance "
    "recommendation. CR-HKGE replaces real interaction supervision with "
    "content-based positive pairs and enriches product embeddings with global "
    "reference semantics through a dedicated cross-reference propagation "
    "mechanism guided by relation-type attention.")
body(doc, "The main contributions of this paper are as follows:")
for txt in [
    "We construct a fragrance-specific heterogeneous KG for the Aromatique "
    "catalog, containing 340 products, 998 entities, 9,250 triples, and seven "
    "relation types.",
    "We introduce a cross-reference semantic propagation layer that transfers "
    "global fragrance reference semantics to local product embeddings through the "
    "inspired_by relation.",
    "We introduce relation-type attention that differentiates the contribution of "
    "heterogeneous fragrance relations during KG propagation.",
    "We construct content-based positive pairs as surrogate supervision for "
    "KGAT-compatible training in a zero historical interaction setting.",
    "We show that CR-HKGE improves over KGAT, BPRMF, CKE, and CFKG on overall "
    "Top-K recommendation metrics, with the largest gains on enriched products.",
]:
    pp = para(doc, before=0, after=2)
    add_run(pp, "•  ", size=10)
    add_run(pp, txt, size=10)

# ================= II. RELATED WORK =================
h1(doc, "II.\tRELATED WORK")
h2(doc, "A.\tCollaborative and KG-Based Recommendation")
body(doc,
    "Matrix factorization with Bayesian personalized ranking (BPRMF) is a widely "
    "used pairwise ranking baseline [3], [20]. To incorporate side information, "
    "collaborative knowledge base embedding (CKE) jointly learns item and KG "
    "embeddings [2], while CFKG casts recommendation as KG completion over a "
    "unified user-item-entity graph [8]. KGAT integrates these ideas through "
    "attentive embedding propagation over a collaborative knowledge graph [1] and "
    "is the direct architectural baseline for this work. Graph collaborative "
    "filtering further confirms the value of high-order connectivity [9]. CR-HKGE "
    "retains the KGAT backbone but adapts it to an interaction-free, "
    "domain-specific fragrance setting.")
h2(doc, "B.\tHeterogeneous KG and Relation Attention")
body(doc,
    "Heterogeneous KG recommenders model multiple node and edge types [10], [11]. "
    "However, several heterogeneous information network models assign uniform "
    "weights to different edge types, which limits recommendation quality [16]. "
    "Heterogeneous KG attention models such as HKGAT extend attentive propagation "
    "but are validated on large datasets and acknowledge the need for cross-domain "
    "semantic bridges [17]. HGKR reports instability on sparser datasets and "
    "suggests stronger KG reasoning as a remedy [14]. CR-HKGE responds to these "
    "observations with learnable relation-type weights and an explicit "
    "cross-reference propagation path.")
h2(doc, "C.\tCold-Start and Conversational Fragrance Recommendation")
body(doc,
    "Cold-start recommendation addresses settings with little or no interaction "
    "data [13]. Empirically, KGAT-style attention has been reported to struggle "
    "with attention allocation on extremely sparse knowledge graphs [15], and "
    "content-derived supervision has been used to compensate for missing feedback "
    "[18], [19]. Conversational recommender systems motivate Top-K retrieval from "
    "natural-language preferences [12]; in this paper, conversational use is the "
    "deployment motivation, while evaluation remains offline. The Aromatique "
    "catalog represents an extreme zero historical interaction case, for which "
    "content-based positive pairs provide surrogate supervision.")

# ================= III. PROPOSED METHOD =================
h1(doc, "III.\tPROPOSED METHOD")
h2(doc, "A.\tProblem Formulation")
body(doc,
    "Let P be the set of perfume products, E the set of entities, R the set of "
    "relation types, and T the set of triples. The fragrance KG is G = (E, R, T), "
    "where each triple (h, r, t) ∈ T links a head entity h to a tail entity t "
    "through relation r. The recommendation task estimates a relevance score")
equation(doc, "ŷ(u, p) = f(eᵤ, eₚ)", 1)
body(doc,
    "between a content/query profile u and a product p. Here u denotes a "
    "content/query profile derived from fragrance attributes, not a real "
    "historical user; no purchase, rating, or click logs are used.")

h2(doc, "B.\tFragrance-Specific Heterogeneous KG Construction")
body(doc,
    "The heterogeneous KG contains seven node types—product, accord, note, "
    "family, global_ref, global_accord, and global_family—and seven relation "
    "types listed in Table I. The inspired_by relation connects a local product "
    "to a global fragrance reference, which in turn carries has_global_accord and "
    "belongs_to_global_family edges. This design exposes a path from a local "
    "product to global semantics, as illustrated in Fig. 2. Products that have an "
    "inspired_by edge are termed enriched (243 products); the remaining products "
    "are standard (97 products).")
figure(doc, os.path.join(ASSETS, "fig2_schema.png"), 3.25,
       "Fig. 2. ",
       "Fragrance-specific heterogeneous KG schema. The inspired_by relation "
       "bridges local products to the global fragrance reference space.")

h2(doc, "C.\tContent-Based Positive-Pair Construction")
body(doc,
    "Because Aromatique has no historical interactions, content-based positive "
    "pairs are generated as surrogate supervision. Each content/query profile "
    "corresponds to one source product, and positive targets are the products "
    "with the highest relevance to that source. The relevance score s(pᵢ, "
    "pⱼ) aggregates weighted Jaccard and indicator terms over local accords, "
    "visual notes, local family, global accords, global family, a "
    "global-to-local accord/family bridge, shared global references, and "
    "sem_similar neighbors, with an additional bonus when both products are "
    "enriched. The eight scoring weights range from 0.20 (enriched bonus) to 3.00 "
    "(same global reference), giving the highest priority to cross-reference "
    "agreement. For each profile, the top-ranked candidates are split into eight "
    "train and four test positive items, yielding 2,720 train and 1,360 test "
    "pairs. These train/test files follow the KGAT format but are content-based "
    "positive pairs, not user behavior.")

h2(doc, "D.\tCR-HKGE Architecture")
body(doc,
    "CR-HKGE preserves the KGAT backbone and adds three novelty components, as "
    "shown in Fig. 3.")
figure(doc, os.path.join(ASSETS, "fig3_architecture.png"), 3.25,
       "Fig. 3. ",
       "CR-HKGE architecture. N1 is the fragrance-specific heterogeneous KG "
       "input; N2 is cross-reference propagation; N3 is relation-type attention.")
body(doc, "1) Layer 1 – TransR KG embedding: "
    "Following KGAT [1] and TransR [4], the plausibility of a triple is")
equation(doc, "g(h, r, t) = ‖ Wᵣ eₕ + eᵣ − Wᵣ eₜ ‖²₂", 2)
body(doc, "and the KG embedding is trained with a pairwise loss over corrupted "
    "tails t′:")
equation(doc, "L_KG = Σ − ln σ( g(h, r, t′) − g(h, r, t) )", 3)
body(doc, "2) Layer 2 – Relation-type attention: "
    "KGAT computes attention π(h, r, t) = (Wᵣ eₜ)ᵀ tanh(Wᵣ "
    "eₕ + eᵣ). CR-HKGE scales this score by a learnable, type-specific "
    "weight:")
equation(doc, "π_CR(h, r, t) = λ̃ᵣ · (Wᵣ eₜ)ᵀ "
              "tanh(Wᵣ eₕ + eᵣ)", 4)
equation(doc, "λ̃ᵣ = |R| · exp(λᵣ) / Σ exp(λᵣ′)", 5)
body(doc,
    "One scalar λᵣ is learned per semantic relation type and shared "
    "between forward and inverse edges. The softmax is rescaled by the number of "
    "relation types so that a uniform initialization preserves the original KGAT "
    "attention scale, while an optional fragrance prior (relation prior strength "
    "1.0) can bias initialization toward informative relations.")
body(doc, "3) Layer 3 – Cross-reference propagation: "
    "For an enriched product p, global reference context is aggregated and "
    "injected into the product embedding:")
equation(doc, "c_ref(p) = α · λ̃_ib · W_cr ( e_g + "
              "Σ π_CR(g, r, t) eₜ )", 6)
body(doc,
    "where g is the global reference reached through inspired_by, λ̃_ib "
    "is its relation weight, and α controls the propagation strength. For "
    "standard products c_ref(p) = 0, enforced by a product mask, so cross-"
    "reference enrichment affects only products that actually have an inspired_by "
    "edge. The final model uses α = 0.1.")
body(doc, "4) Layer 4 – Bi-interaction and prediction: "
    "The cross-reference context is added to the aggregated neighbor message "
    "m_N(p) and combined through the KGAT bi-interaction update")
equation(doc, "eₚ⁽ˡ⁾ = σ( W⁽ˡ⁾ ( "
              "eₚ⁽ˡ⁻¹⁾ + m_N(p) + c_ref(p) ) )", 7)
body(doc, "Layer outputs are concatenated and scored by inner product:")
equation(doc, "eₚ* = eₚ⁽⁰⁾ ∥ … ∥ "
              "eₚ⁽ᴸ⁾ ,    ŷ(u, p) = (eᵤ*)ᵀ eₚ*", 8)

h2(doc, "E.\tTraining Objective")
body(doc,
    "CR-HKGE follows the KGAT alternating schedule. Phase I trains the recommender "
    "with a BPR loss over content-based positive pairs:")
equation(doc, "L_BPR = − Σ ln σ( ŷ(u, i) − ŷ(u, j) )", 9)
body(doc,
    "Phase II trains the TransR/KG embedding and updates the attentive adjacency. "
    "The overall objective is")
equation(doc, "L = L_BPR + L_KG + β ‖ Θ ‖²", 10)

# ================= IV. EXPERIMENTAL SETUP =================
h1(doc, "IV.\tEXPERIMENTAL SETUP")
h2(doc, "A.\tDataset")
body(doc,
    "Experiments use the dataset-aromatique-crhkge-ready dataset, whose key "
    "statistics are summarized in Table I. The dataset contains no historical "
    "purchase, rating, or click logs; the train/test files are content-based "
    "positive pairs used for offline surrogate evaluation in a zero historical "
    "interaction setting. The overall pipeline is shown in Fig. 1, which makes "
    "explicit that no real user interaction history is used.")
figure(doc, os.path.join(ASSETS, "fig1_pipeline.png"), 2.5,
       "Fig. 1. ",
       "Overall research pipeline. No real user interaction history is used; "
       "supervision comes from content-based positive pairs.")

table_caption(doc, "TABLE I", "DATASET AND KNOWLEDGE GRAPH STATISTICS")
make_table(doc,
    ["Statistic / Relation", "Value"],
    [
        ["Products", "340"],
        ["Entities", "998"],
        ["Relation types", "7"],
        ["KG triples", "9,250"],
        ["Enriched products (inspired_by)", "243"],
        ["Standard products", "97"],
        ["Content/query profiles", "340"],
        ["Train positive pairs", "2,720"],
        ["Test positive pairs", "1,360"],
        ["inspired_by", "243"],
        ["has_accord", "1,629"],
        ["has_visual_note", "680"],
        ["belongs_to_family", "340"],
        ["sem_similar", "4,110"],
        ["has_global_accord", "2,017"],
        ["belongs_to_global_family", "231"],
    ],
    aligns=[None, "right"], size=8, col_w=[2.2, 1.0])

h2(doc, "B.\tBaselines")
body(doc,
    "CR-HKGE is compared with four baselines on the same content-based split: "
    "BPRMF, a matrix-factorization ranking baseline [3], [20]; CKE, a "
    "collaborative KG embedding baseline [2]; CFKG, a KG-collaborative baseline "
    "[8]; and KGAT, the main KG-attention architectural baseline [1]. All models "
    "share the same dataset, candidate pool, and evaluation protocol to ensure a "
    "fair comparison.")
h2(doc, "C.\tHyperparameters and Metrics")
body(doc,
    "The embedding and KGE sizes are 64, with propagation layers of size "
    "[64, 32, 16], learning rate 1×10⁻⁴, batch size 64, node "
    "dropout 0.1, message dropout 0.1, and 100 epochs. The CR-HKGE checkpoint is "
    "selected by NDCG@3, the final cross-reference strength is α = 0.1, and "
    "the relation prior strength is 1.0. Recommendation quality is reported with "
    "Precision@K, Recall@K, Hit@K, and NDCG@K for K ∈ {3, 5, 10}, computed "
    "over the full non-training candidate pool. Enriched and standard subsets are "
    "evaluated by restricting the relevant items to products with or without an "
    "inspired_by edge, respectively.")

# ================= V. RESULTS AND DISCUSSION =================
h1(doc, "V.\tRESULTS AND DISCUSSION")
h2(doc, "A.\tOverall Performance")
table_caption(doc, "TABLE II",
              "OVERALL TOP-K PERFORMANCE ON dataset-aromatique-crhkge-ready")
make_table(doc,
    ["Model", "R@3", "P@3", "H@3", "N@3", "R@10", "N@10"],
    [
        ["CFKG", ".0074", ".0098", ".0294", ".0203", ".0309", ".0514"],
        ["BPRMF", ".2110", ".2814", ".6559", ".2809", ".5375", ".5220"],
        ["CKE", ".2265", ".3020", ".6853", ".2987", ".5735", ".5458"],
        ["KGAT", ".2632", ".3510", ".7618", ".3278", ".6912", ".6048"],
        ["CR-HKGE (α=0.1)", ".2949", ".3931", ".8265", ".3663", ".7140", ".6272"],
    ],
    size=8, highlight_last=True,
    col_w=[1.05, 0.4, 0.4, 0.4, 0.4, 0.4, 0.4])
body(doc,
    "Table II reports overall performance. CR-HKGE outperforms all included "
    "baselines on every metric. The improvement over KGAT—the strongest "
    "baseline—confirms that the added cross-reference propagation and "
    "relation-type attention strengthen KGAT-style recommendation. CFKG performs "
    "poorly, indicating that simple collaborative KG modeling is insufficient on "
    "this small, interaction-free catalog, whereas BPRMF and CKE are stronger than "
    "CFKG but remain below the attentive KG models, showing the value of "
    "attentive propagation.")
body(doc,
    "Relative to KGAT, CR-HKGE improves Recall@3 by +0.0316 (0.2632→0.2949), "
    "Precision@3 by +0.0422, Hit@3 by +0.0647, NDCG@3 by +0.0385, Recall@10 by "
    "+0.0228, and NDCG@10 by +0.0223. The largest gains appear at Hit@3 and "
    "Precision@3, indicating that CR-HKGE places relevant candidates earlier in "
    "the recommendation list, which is desirable for a Top-3 conversational "
    "setting.")

h2(doc, "B.\tEnriched vs. Standard Products")
table_caption(doc, "TABLE III",
              "ENRICHED VS. STANDARD PRODUCT PERFORMANCE")
make_table(doc,
    ["Subset / Model", "R@3", "P@3", "H@3", "N@3", "R@10", "N@10"],
    [
        ["Enr. KGAT", ".2598", ".3029", ".6814", ".2993", ".6913", ".5594"],
        ["Enr. CR-HKGE", ".3092", ".3481", ".7640", ".3474", ".7203", ".5912"],
        ["Std. KGAT", ".2921", ".1587", ".4571", ".2219", ".7429", ".4587"],
        ["Std. CR-HKGE", ".2730", ".1492", ".4095", ".2018", ".7413", ".4398"],
    ],
    size=8, col_w=[1.05, 0.4, 0.4, 0.4, 0.4, 0.4, 0.4])
body(doc,
    "Table III separates enriched products (with inspired_by) from standard "
    "products. On enriched products, CR-HKGE improves over KGAT on every metric, "
    "directly supporting the usefulness of cross-reference semantic propagation. "
    "On standard products, where inspired_by is absent and the cross-reference "
    "context is masked to zero, KGAT remains slightly stronger. This behavior is "
    "expected, because cross-reference propagation is designed to benefit "
    "enriched products; CR-HKGE nevertheless remains competitive on standard "
    "products and improves overall recommendation.")

h2(doc, "C.\tSensitivity and Ablation")
body(doc,
    "Varying the cross-reference strength α ∈ {0.25, 0.10, 0.075, 0.05} "
    "shows that α = 0.10 gives the best overall NDCG@3 (0.3663) while "
    "maintaining strong Recall@10 (0.7140). A larger α over-emphasizes "
    "cross-reference paths, whereas a smaller α weakens global reference "
    "enrichment; α = 0.10 provides the best balance for early ranking and "
    "Top-10 retrieval.")
table_caption(doc, "TABLE IV", "ABLATION STUDY (OVERALL)")
make_table(doc,
    ["Variant", "R@3", "P@3", "H@3", "N@3", "R@10", "N@10"],
    [
        ["No novelty modules", ".2625", ".3500", ".7471", ".3244", ".7015", ".6068"],
        ["No cross-reference", ".2824", ".3765", ".7971", ".3496", ".7140", ".6184"],
        ["No fragrance prior", ".2816", ".3755", ".7618", ".3562", ".6993", ".6211"],
        ["Full CR-HKGE (α=0.1)", ".2949", ".3931", ".8265", ".3663", ".7140", ".6272"],
    ],
    size=8, highlight_last=True,
    col_w=[1.05, 0.4, 0.4, 0.4, 0.4, 0.4, 0.4])
body(doc,
    "Table IV reports an ablation in which novelty components are removed. The "
    "full CR-HKGE with α = 0.1 achieves the best overall scores. Removing "
    "the cross-reference path or the fragrance prior reduces early-ranking "
    "quality (NDCG@3 and Hit@3), and disabling all novelty modules reduces the "
    "model to a KGAT-equivalent baseline. These results are consistent with the "
    "overall and subset analyses and indicate that the components are "
    "complementary.")

h2(doc, "D.\tLimitations")
body(doc,
    "Several limitations should be noted. The dataset contains no real historical "
    "user interactions, so positive pairs are surrogate content-based relevance "
    "labels rather than observed behavior, and evaluation is offline rather than "
    "online user testing. The catalog is small (340 products), and the results "
    "should be validated on other fragrance catalogs. Accordingly, the claims are "
    "scoped to improvement over KGAT and the included KG/CF baselines on this "
    "dataset, not to absolute superiority over all possible approaches.")

# ================= VI. CONCLUSION =================
h1(doc, "VI.\tCONCLUSION")
body(doc,
    "This paper proposed CR-HKGE, a cross-reference semantic enrichment model for "
    "fragrance recommendation in a zero historical interaction setting. CR-HKGE "
    "constructs a fragrance-specific heterogeneous KG, propagates global reference "
    "semantics through the inspired_by relation, and applies relation-type "
    "attention to prioritize heterogeneous fragrance relations. Because Aromatique "
    "contains no purchase, rating, or click logs, content-based positive pairs are "
    "used as surrogate supervision for KGAT-compatible training and evaluation. "
    "Experiments show that CR-HKGE improves over KGAT, BPRMF, CKE, and CFKG across "
    "overall Top-K metrics, with the clearest gains on enriched products linked to "
    "global references. Future work will incorporate real user feedback, online "
    "conversational evaluation, larger fragrance catalogs, additional relations "
    "such as season and occasion, and more robust relation-type calibration.")

# ================= REFERENCES =================
h1(doc, "REFERENCES")
refs = [
    "X. Wang, X. He, Y. Cao, M. Liu, and T.-S. Chua, “KGAT: Knowledge Graph "
    "Attention Network for Recommendation,” in Proc. 25th ACM SIGKDD Int. "
    "Conf. Knowledge Discovery & Data Mining (KDD), 2019, pp. 950–958.",
    "F. Zhang, N. J. Yuan, D. Lian, X. Xie, and W.-Y. Ma, “Collaborative "
    "Knowledge Base Embedding for Recommender Systems,” in Proc. 22nd ACM "
    "SIGKDD Int. Conf. Knowledge Discovery & Data Mining (KDD), 2016, "
    "pp. 353–362.",
    "S. Rendle, C. Freudenthaler, Z. Gantner, and L. Schmidt-Thieme, “BPR: "
    "Bayesian Personalized Ranking from Implicit Feedback,” in Proc. 25th "
    "Conf. Uncertainty in Artificial Intelligence (UAI), 2009, pp. 452–461.",
    "Y. Lin, Z. Liu, M. Sun, Y. Liu, and X. Zhu, “Learning Entity and "
    "Relation Embeddings for Knowledge Graph Completion,” in Proc. AAAI Conf. "
    "Artificial Intelligence, 2015, pp. 2181–2187.",
    "A. Bordes, N. Usunier, A. Garcia-Durán, J. Weston, and O. Yakhnenko, "
    "“Translating Embeddings for Modeling Multi-relational Data,” in "
    "Adv. Neural Inf. Process. Syst. (NeurIPS), 2013, pp. 2787–2795.",
    "H. Wang et al., “RippleNet: Propagating User Preferences on the "
    "Knowledge Graph for Recommender Systems,” in Proc. 27th ACM Int. Conf. "
    "Information and Knowledge Management (CIKM), 2018, pp. 417–426.",
    "H. Wang, M. Zhao, X. Xie, W. Li, and M. Guo, “Knowledge Graph "
    "Convolutional Networks for Recommender Systems,” in Proc. World Wide Web "
    "Conf. (WWW), 2019, pp. 3307–3313.",
    "Q. Ai, V. Azizi, X. Chen, and Y. Zhang, “Learning Heterogeneous "
    "Knowledge Base Embeddings for Explainable Recommendation,” Algorithms, "
    "vol. 11, no. 9, p. 137, 2018.",
    "X. Wang, X. He, M. Wang, F. Feng, and T.-S. Chua, “Neural Graph "
    "Collaborative Filtering,” in Proc. 42nd Int. ACM SIGIR Conf. Research "
    "and Development in Information Retrieval (SIGIR), 2019, pp. 165–174.",
    "Q. Guo et al., “A Survey on Knowledge Graph-Based Recommender "
    "Systems,” IEEE Trans. Knowledge and Data Engineering, vol. 34, no. 8, "
    "pp. 3549–3568, 2022.",
    "Z. Wu, S. Pan, F. Chen, G. Long, C. Zhang, and P. S. Yu, “A "
    "Comprehensive Survey on Graph Neural Networks,” IEEE Trans. Neural "
    "Networks and Learning Systems, vol. 32, no. 1, pp. 4–24, 2021.",
    "D. Jannach, A. Manzoor, W. Cai, and L. Chen, “A Survey on Conversational "
    "Recommender Systems,” ACM Computing Surveys, vol. 54, no. 5, "
    "pp. 1–36, 2021.",
    "A. I. Schein, A. Popescul, L. H. Ungar, and D. M. Pennock, “Methods and "
    "Metrics for Cold-Start Recommendations,” in Proc. 25th Int. ACM SIGIR "
    "Conf., 2002, pp. 253–260.",
    "Zhang et al., “HGKR,” Scientific Reports, 2023. [VERIFY metadata].",
    "Rong et al., “ML-KDGATMoco,” Scientific Reports, 2024. "
    "[VERIFY metadata].",
    "Wan et al., “HN-DKG,” Int. J. Intelligent Systems (Wiley), 2024. "
    "[VERIFY metadata].",
    "Zhang et al., “HKGAT,” Applied Intelligence (Springer), 2025. "
    "[VERIFY metadata].",
    "Liang et al., “KGCN-UP,” Scientific Reports, 2025. "
    "[VERIFY metadata].",
    "Ma et al., “KGBPR,” in Proc. Int. Conf. Computing and Pattern "
    "Recognition (ICCPR), 2022. [VERIFY metadata].",
    "Y. Koren, R. Bell, and C. Volinsky, “Matrix Factorization Techniques for "
    "Recommender Systems,” Computer, vol. 42, no. 8, pp. 30–37, 2009.",
]
for i, r in enumerate(refs, 1):
    p = para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    add_run(p, "[%d]  " % i, size=8)
    add_run(p, r, size=8)

doc.save(OUT)
print("saved:", OUT)
